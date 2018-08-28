#!/usr/bin/env python 
# -*- coding:utf-8 -*-
import os
from  docx import Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
import PyQt5.sip
#from PyQt5 import QtGui, QtCore, uic
from PyQt5 import uic
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
import datetime

mainWin, mainClass = uic.loadUiType("src//main.ui")
editWin, editClass = uic.loadUiType("src//edit.ui")
addWin, addClass = uic.loadUiType("src//add.ui")
setWin, setClass = uic.loadUiType("src//setting.ui")
typeWin, typeClass = uic.loadUiType("src//type.ui")
class MyAdd(QWidget, addWin):
    def __init__(self):
        QWidget.__init__(self)
        addWin.__init__(self)
        self.setupUi(self)
        file = open("src//settings.dat", 'r', encoding='UTF-8')
        lines = file.readlines()
        self.label_type_first.setText(lines[1].strip())
        self.add_input1.addItems(lines[2].strip().split(' '))
        self.label_type_second.setText(lines[3].strip())
        self.add_input2.addItems(lines[4].strip().split(' '))
        self.label_type_third.setText(lines[5].strip())
        self.label_type_fourth.setText(lines[6].strip())
        file.close()
        self.btn_cancel.clicked.connect(self.disVis)
        self.btn_save.clicked.connect(self.save)
        self.mymain = None
    def addWin(self, mymain):
        self.mymain = mymain
    def save(self):
        input1 = self.add_input1.currentText()
        input2 = self.add_input2.currentText()
        input3 = self.add_input3.text().strip()
        input4 = self.add_input4.text().strip()
        input5 = self.add_input5.toPlainText()
        if input3 == "" or input4 == "" or input5 == "":
            QMessageBox.warning(self, "输入警告", "输入框不能为空！", QMessageBox.Yes)
            return
        if len(input3) > 20 or len(input4) > 20:
            QMessageBox.warning(self, "输入警告", "手段或类型输入不能过长！", QMessageBox.Yes)
            return
        #print("%s %s %s %s %s" %(input1, input2, input3, input4, input5))
        nowTime = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = input1 + "_" + input2 + "_" + nowTime + ".txt"
        #print(filename)
        file = open("database//" + filename, "w")
        file.write(input3 + "\n")
        file.write(input4 + "\n")
        file.write(input5)
        file.close()
        self.add_input1.setCurrentIndex(0)
        self.add_input2.setCurrentIndex(0)
        self.add_input3.setText("")
        self.add_input4.setText("")
        self.add_input5.setPlainText("")
        self.disVis()
        self.mymain.searchFunc()

    def enVis(self):
        if not self.isVisible():
            self.show()
    def disVis(self):
        if self.isVisible():
            self.hide()

class MyEdit(QWidget, editWin):
    def __init__(self):
        QWidget.__init__(self)
        editWin.__init__(self)
        self.setupUi(self)
        file = open("src//settings.dat", 'r', encoding='UTF-8')
        lines = file.readlines()
        self.label_type_first.setText(lines[1].strip())
        self.edit_input1.addItems(lines[2].strip().split(' '))
        self.label_type_second.setText(lines[3].strip())
        self.edit_input2.addItems(lines[4].strip().split(' '))
        self.label_type_third.setText(lines[5].strip())
        self.label_type_fourth.setText(lines[6].strip())
        file.close()
        self.btn_cancel.clicked.connect(self.disVis)
        self.btn_save.clicked.connect(self.save)
        self.currFileName = ""
        self.mymain = None
    def addWin(self, mymain):
        self.mymain = mymain
    def save(self):
        input1 = self.edit_input1.currentText()
        input2 = self.edit_input2.currentText()
        input3 = self.edit_input3.text().strip()
        input4 = self.edit_input4.text().strip()
        input5 = self.edit_input5.toPlainText()
        if input3 == "" or input4 == "" or input5 == "":
            QMessageBox.warning(self, "输入警告", "输入框不能为空！", QMessageBox.Yes)
            return
        if len(input3) > 10 or len(input4) > 10:
            QMessageBox.warning(self, "输入警告", "手段或类型输入不能过长！", QMessageBox.Yes)
            return
        os.remove("database//" + self.currFileName)
        #print("%s %s %s %s %s" % (input1, input2, input3, input4, input5))
        nowTime = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = input1 + "_" + input2 + "_" + nowTime + ".txt"
        file = open("database//" + filename, "w")
        file.write(input3 + "\n")
        file.write(input4 + '\n')
        file.write(input5)
        file.close()
        self.disVis()
        self.mymain.searchFunc()
    def showCont(self):
        file = open("database//" + self.currFileName, "r")
        content = file.readlines()
        self.edit_input1.setCurrentIndex(self.edit_input1.findText(self.currFileName.split("_")[0]))
        #print((self.currFileName.split("_"))[0])
        #print(self.edit_input1.currentText())
        self.edit_input2.setCurrentIndex(self.edit_input2.findText(self.currFileName.split("_")[1]))
        self.edit_input3.setText(content[0].strip())
        self.edit_input4.setText(content[1].strip())
        details = ""
        for i in range(2, len(content)):
            details += content[i]
        self.edit_input5.setPlainText(details)
        file.close()
    def enVis(self, filename):
        self.currFileName = filename
        if not self.isVisible():
            self.show()
            self.showCont()
    def disVis(self):
        if self.isVisible():
            self.hide()

class MyType(QWidget, typeWin):
    def __init__(self):
        QWidget.__init__(self)
        typeWin.__init__(self)
        self.setupUi(self)
        self.btn_save.clicked.connect(self.disVis)
        self.btn_add_type.clicked.connect(self.addType)
        self.count = 0
        self.text = {}
        self.button = {}
        self.myset = None
        self.type = None
        self.typeNumCount = 0
    def addWin(self, myset):
        self.myset = myset
    def addType(self):
        if self.typeNumCount > 25:
            QMessageBox.warning(self, "输入警告", "类型种类不能过多！", QMessageBox.Yes)
            return
        newType = ""
        newType, ok = QInputDialog.getText(self, "输入框", "输入新类型名称:", QLineEdit.Normal)
        if ok:
            if (len(newType.strip()) == 0):
                QMessageBox.warning(self, "输入警告", "类型名称不能为空！", QMessageBox.Yes)
                return
            if len(newType.strip()) > 7:
                QMessageBox.warning(self, "输入警告", "类型名称不能过长！", QMessageBox.Yes)
                return
            if self.type.findText(newType.strip()) >= 0:
                QMessageBox.warning(self, "输入警告", "该类型已存在！", QMessageBox.Yes)
                return
            self.type.addItem(newType.strip())
            self.typeNumCount += 1
        self.showCont()
    def showCont(self):
        if self.count > 0:
           self.area_type.takeWidget()

        widget = QWidget();
        widget.setMinimumSize(10, 10);
        widget.setMaximumSize(10, 10);

        GLayout = QGridLayout()
        GLayout.setSizeConstraint(QGridLayout.SetMinAndMaxSize)

        self.text = {}
        self.button = {}
        self.count = 1

        typeNum = self.type.count()
        self.typeNumCount = typeNum
        for i in range(typeNum):
            self.text[i] = QPushButton(self.type.itemText(i))
            self.button[i] = QPushButton("删除")
            self.text[i].setMinimumSize(160, 28);
            self.text[i].setMaximumSize(160, 28);
            self.text[i].setStyleSheet("font:13pt '.SF NS Text'; background-color:rgba(255, 255, 255, 0);");
            self.button[i].setMinimumSize(75,28);
            self.button[i].setMaximumSize(75, 28);
            self.button[i].setStyleSheet("font:13pt '.SF NS Text'; background-color:rgba(255, 255, 255, 0);text-decoration:underline;");
            self.button[i].setCursor(QCursor(Qt.PointingHandCursor));
            self.button[i].setObjectName(str(i))
            self.button[i].clicked.connect(self.delType)
            GLayout.addWidget(self.text[i], i, 0)
            GLayout.addWidget(self.button[i], i, 1)

        widget.setLayout(GLayout)
        self.area_type.setWidget(widget);
    def delType(self):
        btn_id = self.sender().objectName()
        self.type.removeItem(int(btn_id))
        self.typeNumCount -= 1
        self.showCont()
    def enVis(self, type_id):
        if type_id == 1:
            self.type = self.myset.type_first
        if type_id == 2:
            self.type = self.myset.type_second
        if not self.isVisible():
            self.show()
            self.showCont()
    def disVis(self):
        if self.isVisible():
            self.hide()

class MySet(QWidget, setWin):
    def __init__(self):
        QWidget.__init__(self)
        setWin.__init__(self)
        self.setupUi(self)
        self.btn_cancel.clicked.connect(self.disVis)
        self.btn_save.clicked.connect(self.save)
        self.btn_type_first.clicked.connect(self.editType)
        self.btn_type_second.clicked.connect(self.editType)
        self.mytype = None
        self.mymain = None

    def addWin(self, mytype, mymain):
        self.mytype = mytype
        self.mymain = mymain
    def editType(self):
        btn_id = self.sender().objectName()
        if btn_id == "btn_type_first":
            self.mytype.enVis(1)
        if btn_id == "btn_type_second":
            self.mytype.enVis(2)
    def save(self):
        setStr = []
        input = self.title.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "系统标题不能为空！", QMessageBox.Yes)
            return
        if len(input) > 20:
            QMessageBox.warning(self, "输入警告", "系统标题不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = self.label_type_first.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "搜索条件标签一不能为空！", QMessageBox.Yes)
            return
        if len(input) > 8:
            QMessageBox.warning(self, "输入警告", "搜索条件标签一不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = ""
        typeNum = self.type_first.count()
        for i in range(typeNum):
            input += self.type_first.itemText(i) + ' '
        if input.strip() == "":
            QMessageBox.warning(self, "输入警告", "类型种类不能为空！", QMessageBox.Yes)
            return
        setStr.append(input.rstrip())
        input = self.label_type_second.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "搜索条件标签二不能为空！", QMessageBox.Yes)
            return
        if len(input) > 8:
            QMessageBox.warning(self, "输入警告", "搜索条件标签二不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = ""
        typeNum = self.type_second.count()
        for i in range(typeNum):
            input += self.type_second.itemText(i) + ' '
        if input.strip() == "":
            QMessageBox.warning(self, "输入警告", "类型种类不能为空！", QMessageBox.Yes)
            return
        setStr.append(input.rstrip())
        input = self.label_type_third.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "结果类型标签一不能为空！", QMessageBox.Yes)
            return
        if len(input) > 8:
            QMessageBox.warning(self, "输入警告", "结果类型标签一不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = self.label_type_fourth.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "结果类型标签二不能为空！", QMessageBox.Yes)
            return
        if len(input) > 8:
            QMessageBox.warning(self, "输入警告", "结果类型标签二不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)

        ##function area
        input = self.func_add.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "添加功能标签不能为空！", QMessageBox.Yes)
            return
        if len(input) > 7:
            QMessageBox.warning(self, "输入警告", "添加功能标签不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = self.func_edit.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "修改功能标签不能为空！", QMessageBox.Yes)
            return
        if len(input) > 7:
            QMessageBox.warning(self, "输入警告", "修改功能标签不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = self.func_del.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "删除功能标签不能为空！", QMessageBox.Yes)
            return
        if len(input) > 7:
            QMessageBox.warning(self, "输入警告", "删除功能标签不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)
        input = self.func_save.text().strip()
        if input == "":
            QMessageBox.warning(self, "输入警告", "保存功能标签不能为空！", QMessageBox.Yes)
            return
        if len(input) > 7:
            QMessageBox.warning(self, "输入警告", "保存功能标签不能过长！", QMessageBox.Yes)
            return
        setStr.append(input)

        res = QMessageBox.warning(self, "", "确认是否保存系统配置？", QMessageBox.No | QMessageBox.Yes)
        if res == QMessageBox.No:
            return
        file = open("src//settings.dat", "w", encoding='UTF-8')
        for i in range(len(setStr)):
            file.write(setStr[i] + '\n')
        file.close()
        QMessageBox.information(self, "", "系统配置保存成功！重启软件生效！", QMessageBox.Yes)
        self.disVis()
    def showCont(self):
        self.type_first.clear()
        self.type_second.clear()
        file = open("src//settings.dat", 'r', encoding='UTF-8')
        lines = file.readlines()
        self.title.setText(lines[0].strip())
        self.label_type_first.setText(lines[1].strip())
        self.type_first.addItems(lines[2].strip().split(' '))
        self.label_type_second.setText(lines[3].strip())
        self.type_second.addItems(lines[4].strip().split(' '))
        self.label_type_third.setText(lines[5].strip())
        self.label_type_fourth.setText(lines[6].strip())
        self.func_add.setText(lines[7].strip())
        self.func_edit.setText(lines[8].strip())
        self.func_del.setText(lines[9].strip())
        self.func_save.setText(lines[10].strip())
        file.close()
    def enVis(self):
        if not self.isVisible():
            self.show()
            self.showCont()
            self.mymain.hide()
    def disVis(self):
        if self.isVisible():
            self.mymain.show()
            self.hide()

class MyApp(QMainWindow, mainWin):
    def __init__(self):
        QMainWindow.__init__(self)
        mainWin.__init__(self)
        self.setupUi(self)
        self.inits()
        self.setWindowIcon(QIcon("src//search.ico"))
        self.area_brief.setSizePolicy(QSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding))
        self.area_brief.setWidgetResizable(True)
        self.btn_search.clicked.connect(self.searchFunc)
        self.func_exit.clicked.connect(self.exit)
        self.func_add.clicked.connect(self.add)
        self.func_edit.clicked.connect(self.edit)
        self.func_del.clicked.connect(self.delete)
        self.func_save.clicked.connect(self.save)
        self.actionsettings.triggered.connect(self.setting)

        self.filelist = []
        self.result = []
        self.currfileid = -1
        self.btn_id = -1
        self.text0 = {}
        self.text1 = {}
        self.text2 = {}
        self.button = {}
        self.count = 0
        self.myadd = None
        self.myedit = None
        self.myset = None
    def setting(self):
        print("settings!")
        self.myset.enVis()
    def inits(self):
        file = open("src//settings.dat", 'r', encoding='UTF-8')
        lines = file.readlines()
        self.title.setText(lines[0].strip())
        self.label_type_first.setText(lines[1].strip())
        self.type_first.addItems(lines[2].strip().split(' '))
        self.label_type_second.setText(lines[3].strip())
        self.type_second.addItems(lines[4].strip().split(' '))
        self.label_type_third.setText(lines[5].strip())
        self.label_type_fourth.setText(lines[6].strip())
        self.func_add.setText(lines[7].strip())
        self.func_edit.setText(lines[8].strip())
        self.func_del.setText(lines[9].strip())
        self.func_save.setText(lines[10].strip())
        file.close()
    def save(self):
        self.checkId()
        if self.currfileid == -1:
            QMessageBox.warning(self, "", "请选择保存对象！", QMessageBox.Yes)
            return
        file = open("database//" + self.filelist[self.currfileid], "r")
        content = file.readlines()
        if len(content) < 3:
            QMessageBox.warning(self, "", "文件内容缺失！", QMessageBox.Yes)
            return
        file.close()
        directory = QFileDialog.getExistingDirectory(self, "选取文件夹", "C:/")  # 起始路径
        if directory == "":
            return
        nowTime = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = self.type_first.currentText() + "_" + self.type_second.currentText() + "_" + nowTime + ".docx"
        dirList = directory.split("/")
        dir = ""
        for i in range(len(dirList)):
            dir += dirList[i] + "//"

        document = Document(docx=os.path.join(os.getcwd(), 'src', 'default.docx'))
        paragraph = document.add_paragraph()
        run1 = paragraph.add_run(u''+self.type_first.currentText() + "  " + self.type_second.currentText())
        run1.font.size = Pt(16)
        run1.font.name = u'仿宋'
        run1.bold = True
        r = run1._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

        for i in range(len(content)):
            paragraph = document.add_paragraph()
            # 设置字号
            run = paragraph.add_run(u'' + content[i].strip())
            run.font.size = Pt(14)
            if i < 2:
                run.bold = True
            # 设置字体
            run.font.name = u'仿宋'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        document.save(dir + filename)
        QMessageBox.information(self, "", "文件保存成功！", QMessageBox.Yes)


    def delete(self):
        self.checkId()
        if self.currfileid == -1:
            QMessageBox.warning(self, "", "请选择删除对象！", QMessageBox.Yes)
            return
        res = QMessageBox.warning(self, "", "确认是否删除？", QMessageBox.No | QMessageBox.Yes)
        if res == QMessageBox.Yes:
            os.remove("database//" + self.filelist[self.currfileid])
            self.searchFunc()
    def addWin(self, myadd, myedit, myset):
        self.myadd = myadd
        self.myedit = myedit
        self.myset = myset
    def add(self):
        self.myadd.enVis()
    def edit(self):
        self.checkId()
        if self.currfileid == -1:
            QMessageBox.warning(self, "", "请选择编辑对象！", QMessageBox.Yes)
            return
        self.myedit.enVis(self.filelist[self.currfileid])
    def searchFunc(self):
        self.currfileid = -1
        self.btn_id = -1
        self.area_detail.setText("")
        self.label_detail.show()
        type_one = self.type_first.currentText()
        type_two = self.type_second.currentText()
        print("search: %s %s" % (type_one, type_two))
        filenamepre = type_one + "_" + type_two
        self.filelist = []
        self.result = []
        filenames  = os.listdir("database")
        matchnum = 0
        for i in range(len(filenames)):
            filepath = os.path.join("database", filenames[i])
            if not os.path.isfile(filepath):
                continue
            if filenames[i].startswith(filenamepre):
                self.filelist.append(filenames[i])
                file = open(filepath, "r")
                content = file.readlines()
                self.result.append([])
                self.result[matchnum].append(content[0].strip())
                self.result[matchnum].append(content[1].strip())
                details = ""
                for j in range(2, len(content)):
                    details += content[j]
                self.result[matchnum].append(details)
                matchnum += 1
                file.close()

        self.showList(self.result)


    def showList(self, result):
        if self.count > 0:
           self.area_brief.takeWidget()

        widget = QWidget();
        widget.setMinimumSize(10, 10);
        widget.setMaximumSize(10, 10);

        if result == []:
            self.text0 = {}
            self.text1 = {}
            self.text2 = {}
            self.button = {}
            self.count = 1
            noresult = QLabel("无结果")
            noresult.setStyleSheet("font:30pt '.SF NS Text'; color:rgb(204, 204, 204);")
            GLayout = QVBoxLayout()
            GLayout.setSizeConstraint(QVBoxLayout.SetMinAndMaxSize)
            GLayout.addWidget(noresult)
            widget.setLayout(GLayout)
            self.area_brief.setWidget(widget)
            return

        GLayout = QGridLayout()
        GLayout.setSizeConstraint(QGridLayout.SetMinAndMaxSize)

        self.text0 = {}
        self.text1 = {}
        self.text2 = {}
        self.button = {}
        self.count = 1
        for i in range(len(result)):
            self.text0[i] = QRadioButton(str(i + 1))
            self.text1[i] = QPushButton(result[i][0])
            self.text2[i] = QPushButton(result[i][1])
            self.button[i] = QPushButton("查看详情")
            self.text0[i].setMinimumSize(48, 28);
            self.text0[i].setMaximumSize(48, 28);
            #self.text0[i].setAlignment(Qt.AlignCenter);
            self.text0[i].setStyleSheet("font:13pt '.SF NS Text';");
            self.text0[i].setCursor(QCursor(Qt.PointingHandCursor));
            self.text1[i].setMinimumSize(174, 28);
            self.text1[i].setMaximumSize(174, 28);
            self.text1[i].setStyleSheet("font:13pt '.SF NS Text'; background-color:rgba(255, 255, 255, 0);");
            self.text1[i].setCursor(QCursor(Qt.PointingHandCursor));
            self.text1[i].setObjectName("text1_" + str(i))
            self.text1[i].clicked.connect(self.setSelect)
            self.text2[i].setMinimumSize(178, 28);
            self.text2[i].setMaximumSize(178, 28);
            self.text2[i].setStyleSheet("font:13pt '.SF NS Text'; background-color:rgba(255, 255, 255, 0);");
            self.text2[i].setCursor(QCursor(Qt.PointingHandCursor));
            self.text2[i].setObjectName("text2_ " + str(i))
            self.text2[i].clicked.connect(self.setSelect)
            self.button[i].setMinimumSize(86,28);
            self.button[i].setMaximumSize(86, 28);
            self.button[i].setStyleSheet("font:13pt '.SF NS Text'; background-color:rgba(255, 255, 255, 0);text-decoration:underline;");
            self.button[i].setCursor(QCursor(Qt.PointingHandCursor));
            self.button[i].setObjectName(str(i))
            self.button[i].clicked.connect(self.showDetails)
            GLayout.addWidget(self.text0[i], i, 0)
            GLayout.addWidget(self.text1[i], i, 1)
            GLayout.addWidget(self.text2[i], i, 2)
            GLayout.addWidget(self.button[i], i, 3)

        widget.setLayout(GLayout)
        self.area_brief.setWidget(widget);

    def setSelect(self):
        btn_id = self.sender().objectName()
        self.btn_id = int(btn_id.split("_")[1])
        self.text0[self.btn_id].setChecked(True)
    def showDetails(self):
        self.label_detail.hide()
        btn_id = self.sender().objectName()
        self.btn_id = int(btn_id)
        self.text0[int(btn_id)].setChecked(True)
        self.area_detail.setText(self.result[int(btn_id)][2])

    def exit(self):
        self.close()

    def clearLayout(self, layout):
        while layout.count():
            child = layout.takeAt(0)
            if child.widget() is not None:
                child.widget().deleteLater()
            elif child.layout() is not None:
                self.clearLayout(child.layout())
    def checkId(self):
        for key in self.text0.keys():
            if self.text0[key].isChecked():
                self.currfileid = key
                return
        self.currfileid = -1
